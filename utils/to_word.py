from docx import Document
import io
import base64

from mosreg.models import Payment


def generate_document(item_name: Payment) -> io.BytesIO:
    document = Document()
    document.add_heading(f"Document for {item_name.num}", level=1)
    content = f"This is the content of the document for {item_name.num}."
    document.add_paragraph(content)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    binary_content = buffer.read()
    base64_encoded = base64.b64encode(binary_content)
    print(base64_encoded, 838383)
    return base64_encoded
